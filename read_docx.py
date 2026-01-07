#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
读取.docx文件内容
"""

from docx import Document
import sys

def read_docx(file_path):
    """读取docx文件内容"""
    try:
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # 也读取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)
    except Exception as e:
        return f"读取失败: {str(e)}"

if __name__ == "__main__":
    file_path = "assets/自动化交易程序工作原理_ETH5m_Fakeout.docx"
    content = read_docx(file_path)
    print(content)
